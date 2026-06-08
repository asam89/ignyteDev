"""
Document Parser — extracts text from uploaded files.

Supports: PDF (PyPDF2), Word (.docx), plain text, code files, CSV,
and images (passed through to LLM vision API).
"""

import io
import logging
from pathlib import Path

log = logging.getLogger(__name__)

IMAGE_EXTENSIONS = {
    ".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg", ".ico",
}

TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".rst", ".csv", ".tsv",
    ".py", ".js", ".ts", ".tsx", ".jsx", ".html", ".css", ".scss",
    ".yml", ".yaml", ".json", ".sh", ".toml", ".cfg", ".ini",
    ".xml", ".sql", ".rb", ".go", ".rs", ".java", ".kt", ".swift",
    ".c", ".cpp", ".h", ".hpp", ".vue", ".svelte",
}


def is_image(filename: str) -> bool:
    return Path(filename).suffix.lower() in IMAGE_EXTENSIONS


def extract_text(content: bytes, filename: str) -> str:
    """Extract readable text from an uploaded file."""
    lower = filename.lower()
    suffix = Path(lower).suffix

    if suffix in TEXT_EXTENSIONS:
        return content.decode("utf-8", errors="replace")

    if suffix == ".pdf":
        return _extract_pdf(content, filename)

    if suffix in (".docx",):
        return _extract_docx(content, filename)

    if suffix == ".doc":
        return "(Legacy .doc format — please convert to .docx or PDF)"

    if suffix in IMAGE_EXTENSIONS:
        return ""  # Images are handled via vision API, not text

    return f"(Binary file: {filename})"


def _extract_pdf(content: bytes, filename: str) -> str:
    """Extract text from PDF using PyPDF2."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(content))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"--- Page {i + 1} ---\n{text.strip()}")
        if pages:
            return "\n\n".join(pages)
        return "(PDF — no extractable text, may be image-based)"
    except ImportError:
        log.warning("PyPDF2 not installed, falling back to basic PDF extraction")
        return _extract_pdf_basic(content)
    except Exception as e:
        log.warning(f"PyPDF2 extraction failed for {filename}: {e}")
        return _extract_pdf_basic(content)


def _extract_pdf_basic(content: bytes) -> str:
    """Fallback PDF extraction."""
    try:
        text = content.decode("latin-1", errors="replace")
        readable = "".join(c for c in text if c.isprintable() or c in "\n\r\t")
        return readable[:5000] if readable.strip() else "(PDF — could not extract text)"
    except Exception:
        return "(PDF — could not extract text)"


def _extract_docx(content: bytes, filename: str) -> str:
    """Extract text from Word .docx files."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        if paragraphs:
            return "\n\n".join(paragraphs)
        return f"(Word document {filename} — no extractable text)"
    except ImportError:
        log.warning("python-docx not installed, cannot extract .docx")
        return "(Word .docx — install python-docx for extraction)"
    except Exception as e:
        log.warning(f"DOCX extraction failed for {filename}: {e}")
        return f"(Word document extraction failed: {e})"
